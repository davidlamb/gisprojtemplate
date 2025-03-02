#Selenium for the scraping
from selenium import webdriver

from bs4 import BeautifulSoup


from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor, Inches
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path
import xml.etree.ElementTree as ET
from datetime import datetime,timezone
import re

def looptext(docelement, htmlelement, headingText = False):
    for x in htmlelement.find_all(string=True):
        if headingText is None:
            headingText = ""
        parents = x.findParents()
        italic = False
        bold = False
        for par in parents:
            if par.name == "em":
                italic = True
            if par.name == "b":
                bold = True

        if x.parent.name != "span":
            segment = docelement.add_run(x)
            segment.font.italic = italic
            segment.font.bold = bold               
        elif x.parent.name == "span":
            #print(f"span {x}")
            #class_name = x.parent.find(attrs={"name": "class"})
            span_tag = x.parent
            span_par = span_tag.parent
            skip = False
            if headingText is False:
                if span_par.has_attr('class'):
                    if "paragraph-hierarchy" in span_par['class'] or "paren" in span_par['class']:
                        skip = True
            else:
                if "paragraph-hierarchy" in span_tag['class'] or "paren" in span_tag['class']:
                    segment = docelement.add_run(x)
                    segment.font.italic = italic
                    segment.font.bold = bold       

            if skip is False:
                if "diff-html-removed" in span_tag['class']:
                    segment = docelement.add_run(x)
                    segment.font.strike = True
                    segment.font.color.rgb = RGBColor(255, 0, 0)
                    segment.font.italic = italic
                    segment.font.bold = bold    
                elif "diff-html-added" in span_tag['class']:
                    segment = docelement.add_run(x)
                    segment.font.underline = True
                    segment.font.color.rgb = RGBColor(255, 0, 0)
                    segment.font.italic = italic
                    segment.font.bold = bold    
                elif "diff-html-changed" in span_tag['class']:
                    segment = docelement.add_run(x)
                    segment.font.color.rgb = RGBColor(20, 255, 0)
                    segment.font.italic = italic
                    segment.font.bold = bold       

def loop_children(htmlelement,docobj:Document,styleforparagraph = None):
    
    for chld in htmlelement:
        if chld.name == "p":
            
            if chld.has_attr('class'):
                if styleforparagraph is None:
                    p = docobj.add_paragraph()
                else:
                    p = docobj.add_paragraph(style=styleforparagraph)
                if "indent-1" in chld['class']:
                    p.paragraph_format.left_indent = Inches(0.25)
                if "indent-2" in chld['class']:
                    p.paragraph_format.left_indent = Inches(0.35)
                if "indent-3" in chld['class']:
                    p.paragraph_format.left_indent = Inches(0.5)
                if "indent-4" in chld['class']:
                    p.paragraph_format.left_indent = Inches(0.6)
                if "indent-5" in chld['class']:
                    p.paragraph_format.left_indent = Inches(0.6)
                ph = chld.findChild("span", {"class": "paragraph-hierarchy"})
                if ph is not None:
                    looptext(p,ph,True)
                looptext(p,chld,False)
            else:
                if styleforparagraph is None:
                    p = docobj.add_paragraph()
                else:
                    p = docobj.add_paragraph(style=styleforparagraph)
                looptext(p, chld)

        elif chld.name == "span":
            if chld.has_attr('class'):
                if "diff-html-removed" in chld['class'] or "diff-html-added" in chld['class']:# or "paragraph-hierarchy" in chld['class']:
                #if "paragraph-hierarchy" in chld['class']:
                    if styleforparagraph is None:
                        p = docobj.add_paragraph()
                    else:
                        p = docobj.add_paragraph(style=styleforparagraph)
                    looptext(p, chld)
                    #pass
        elif chld.name == "div":
            loop_children(chld.findChildren(recursive=False),docobj,styleforparagraph)

pth = Path("C:/Temp")
fn = pth / "trackchanges_blank.docx"
doc = Document(fn)
#np = doc.add_paragraph()
obj_styles = doc.styles
obj_charstyle = obj_styles.add_style('InsertedParagraph', WD_STYLE_TYPE.PARAGRAPH)
obj_font = obj_charstyle.font
obj_font.underline = True
obj_font.color.rgb = RGBColor(255, 0, 0)



driver = webdriver.Chrome(r"c:\working\chromedriver.exe")
driver.implicitly_wait(30)
url = "https://www.ecfr.gov/compare/2024-06-03/to/2024-06-02/title-49/subtitle-A/part-24"
driver.get(url)

html = driver.page_source

soup = BeautifulSoup(html,features='lxml')

content_block = soup.find(id='part-24') #driver.find_element_by_id('part-24')
title = content_block.findChild('h1')
heading = doc.add_heading()

looptext(heading,title)


authority = content_block.findChild("div", {"class": "authority"})
p = doc.add_paragraph()

authority_heading = p.add_run("Authority: ")
authority_heading.font.bold = True

authority_paragraph = authority.findChild("p")
looptext(p,authority_paragraph)


main_element = content_block.findChild("div", {"class": "source"})
p = doc.add_paragraph()
p_segment = p.add_run("Source: ")
p_segment.font.bold = True

main_element_paragraph = main_element.findChild("p")
looptext(p,main_element_paragraph)


all_subparts = content_block.findChildren("div", {"class": "subpart"})

for subpart in all_subparts:
    subpart_heading = subpart.findChild("h2")
    h = doc.add_heading(level=2)
    looptext(h,subpart_heading)
    all_sections = content_block.findChildren("div", {"class": "section"})
    for section in all_sections:
        if section.has_attr('class'):
            section_parent = section.parent
            styleFormat = None
            if section_parent.has_attr('class'):
                if 'diff-html-added' in section_parent['class']:
                    styleFormat = "InsertedParagraph"

            if 'section' in section['class']:
                section_heading = section.findChild("h4")
                h = doc.add_heading(level=3)
                looptext(h,section_heading)
                section_children = section.findChildren(recursive=False)
                loop_children(section_children,doc, styleFormat)

            doc.save(r"c:\temp\test.docx")
    break
doc.save(r"c:\temp\test.docx")
#title = content_block.find_element_by_tag_name("h1")
#print(title.getText())
#spans = title.find_elements_by_tag_name("span")
#bodytext = title.get_attribute('innerHTML')
#title.find_all(text=True)
#bodytext_groups = []
#if len(spans)>0:
#    for span in spans:
#        if span.get_attribute("class") == "diff-html-removed":
#            print(span.text)
            
#print(title.get_attribute('innerHTML'))
#for x in title.find_elements_by_css_selector("*"):
#    print(x.get_attribute('innerHTML'))